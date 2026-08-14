#!/usr/bin/env python3
"""文本提取：docx（python-docx）/ xlsx（zipfile+ElementTree，零第三方依赖）。
用法: extract.py docx <file> | extract.py xlsx <file>
"""
import re
import sys
import zipfile
from xml.etree import ElementTree as ET

SS_NS = "{http://schemas.openxmlformats.org/spreadsheetml/2006/main}"


def extract_docx(path):
    from docx import Document

    doc = Document(path)
    out = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            out.append(t)
    for tbl in doc.tables:
        for row in tbl.rows:
            out.append(" | ".join(c.text.strip() for c in row.cells))
    return "\n".join(out)


def extract_xlsx(path):
    z = zipfile.ZipFile(path)
    shared = []
    try:
        ss = ET.fromstring(z.read("xl/sharedStrings.xml"))
        for si in ss:
            shared.append(
                "".join(t.text or "" for t in si.iter(SS_NS + "t"))
            )
    except KeyError:
        pass
    out = []
    for name in z.namelist():
        m = re.match(r"xl/worksheets/sheet(\d+)\.xml", name)
        if not m:
            continue
        root = ET.fromstring(z.read(name))
        rows = []
        for row in root.iter(SS_NS + "row"):
            cells = []
            for c in row.iter(SS_NS + "c"):
                v = c.find(SS_NS + "v")
                if v is None:
                    cells.append("")
                    continue
                idx = int(v.text or 0)
                if c.get("t") == "s" and idx < len(shared):
                    cells.append(shared[idx])
                else:
                    cells.append(v.text or "")
            rows.append(" | ".join(cells))
        out.append(f"【工作表 {m.group(1)}】")
        out.extend(rows)
    return "\n".join(out)


def main():
    kind, path = sys.argv[1], sys.argv[2]
    if kind == "docx":
        print(extract_docx(path))
    elif kind == "xlsx":
        print(extract_xlsx(path))
    else:
        sys.exit(f"unsupported kind: {kind}")


if __name__ == "__main__":
    main()
