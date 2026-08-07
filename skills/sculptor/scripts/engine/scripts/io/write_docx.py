#!/usr/bin/env python3
"""markdown → docx。用法:
  write_docx.py <in.md> <out.docx> [--official] [--redhead] [--academic]
--official: GB/T 9704-2012 公文排版（A4，3号仿宋正文，黑体/楷体层级标题，页码）
--redhead:  红头文件（红色发文机关标志 + 红色分隔线 + 发文字号；需 --official）
--academic: 学术论文排版（宋体小四正文 1.5 倍行距、黑体标题、楷体二级标题）
markdown 首部支持 front-matter 提供发文字号：
  ---
  文号: 国办发〔2026〕12号
  ---
"""
import re
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

RED = RGBColor(0xFF, 0x00, 0x00)


def set_font(run, east, size_pt, bold=False, color=None):
    run.font.name = "Times New Roman"
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), east)
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_page_number(doc):
    footer = doc.sections[0].footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run("— ")
    set_font(r1, "宋体", 14)
    rn = p.add_run()
    set_font(rn, "宋体", 14)
    f1 = OxmlElement("w:fldChar")
    f1.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = " PAGE "
    f2 = OxmlElement("w:fldChar")
    f2.set(qn("w:fldCharType"), "end")
    rn._r.append(f1)
    rn._r.append(it)
    rn._r.append(f2)
    r2 = p.add_run(" —")
    set_font(r2, "宋体", 14)


def add_red_line(doc):
    p = doc.add_paragraph()
    ppr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "FF0000")
    pbdr.append(bottom)
    ppr.append(pbdr)
    return p


def parse_front_matter(lines):
    if not lines or lines[0].strip() != "---":
        return {}, lines
    fm = {}
    i = 1
    while i < len(lines) and lines[i].strip() != "---":
        m = re.match(r"^([^:：]+)[:：]\s*(.+)$", lines[i].strip())
        if m:
            fm[m.group(1).strip()] = m.group(2).strip()
        i += 1
    if i < len(lines):
        i += 1
    return fm, lines[i:]


def official_paragraph(doc, line, is_title):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = Pt(28)
    pf.space_after = Pt(0)
    if is_title:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.add_run(line), "方正小标宋简体", 22)
        return
    if re.match(r"^[一二三四五六七八九十]+、", line):
        set_font(p.add_run(line), "黑体", 16)
        return
    if re.match(r"^（[一二三四五六七八九十]+）", line):
        pf.first_line_indent = Pt(32)
        set_font(p.add_run(line), "楷体_GB2312", 16)
        return
    pf.first_line_indent = Pt(32)
    set_font(p.add_run(line), "仿宋_GB2312", 16)


def official_doc(md, out, redhead):
    lines = [l.rstrip("\n") for l in open(md, encoding="utf-8")]
    fm, lines = parse_front_matter(lines)
    doc = Document()
    sec = doc.sections[0]
    sec.page_height = Mm(297)
    sec.page_width = Mm(210)
    sec.top_margin = Mm(37)
    sec.bottom_margin = Mm(35)
    sec.left_margin = Mm(28)
    sec.right_margin = Mm(26)
    if redhead:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.add_run("××× 文件"), "方正小标宋简体", 22, color=RED)
        if fm.get("文号"):
            p2 = doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_font(p2.add_run(fm["文号"]), "仿宋_GB2312", 16)
        add_red_line(doc)
    title_done = False
    body = []
    for line in lines:
        if not line.strip():
            continue
        h = re.match(r"^#\s+(.+)$", line)
        if h:
            if not title_done:
                official_paragraph(doc, h.group(1), True)
                title_done = True
            else:
                official_paragraph(doc, h.group(1), False)
            continue
        body.append(line)
    for line in body:
        official_paragraph(doc, line, False)
    # 落款与成文日期：最后两个非空段落右对齐（发文机关署名 + 日期）
    if len(body) >= 2:
        for p in doc.paragraphs[-2:]:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_number(doc)
    doc.save(out)
    print(out)


def plain_doc(md, out):
    doc = Document()
    for raw in open(md, encoding="utf-8"):
        line = raw.rstrip("\n")
        if not line.strip():
            continue
        h = re.match(r"^(#{1,6})\s+(.+)$", line)
        if h:
            doc.add_heading(h.group(2), level=min(len(h.group(1)), 4))
        elif line.lstrip().startswith(("- ", "* ")):
            doc.add_paragraph(line.lstrip()[2:], style="List Bullet")
        elif line.startswith("> "):
            doc.add_paragraph(line[2:], style="Intense Quote")
        else:
            doc.add_paragraph(line)
    doc.save(out)
    print(out)


def academic_doc(md, out):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Mm(25)
    sec.left_margin = sec.right_margin = Mm(25)
    for raw in open(md, encoding="utf-8"):
        line = raw.rstrip("\n")
        if not line.strip():
            continue
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing = 1.5
        pf.space_after = Pt(0)
        h = re.match(r"^#\s+(.+)$", line)
        if h:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_font(p.add_run(h.group(1)), "黑体", 16)
            continue
        h2 = re.match(r"^##\s+(.+)$", line)
        if h2:
            set_font(p.add_run(h2.group(1)), "黑体", 12)
            continue
        if re.match(r"^[一二三四五六七八九十]+、", line):
            set_font(p.add_run(line), "黑体", 12)
            continue
        if re.match(r"^（[一二三四五六七八九十]+）", line):
            pf.first_line_indent = Pt(24)
            set_font(p.add_run(line), "楷体_GB2312", 12)
            continue
        pf.first_line_indent = Pt(24)
        set_font(p.add_run(line), "宋体", 12)
    doc.save(out)
    print(out)


def main():
    args = sys.argv[1:]
    official = "--official" in args
    academic = "--academic" in args
    redhead = "--redhead" in args
    args = [a for a in args if a not in ("--official", "--redhead", "--academic")]
    md, out = args[0], args[1]
    if official:
        official_doc(md, out, redhead)
    elif academic:
        academic_doc(md, out)
    else:
        plain_doc(md, out)


if __name__ == "__main__":
    main()
