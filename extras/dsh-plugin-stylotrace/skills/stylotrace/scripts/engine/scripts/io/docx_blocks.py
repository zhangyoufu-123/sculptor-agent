#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""DOCX 块级提取 / 回填（run 级格式保留）v0.56。
用法：
  python3 docx_blocks.py extract <in.docx>              # → JSON {blocks:[{id,type,text}]}
  python3 docx_blocks.py apply <in.docx> <repl.json> <out.docx>  # 按 id 回填，保留原 run 格式

块 id：P<i>（正文段落，文档顺序）；T<ti>_R<ri>_C<ci>_P<pi>（表格单元格段落）。
回填策略：新文本写入段落首个 run（保留其字体/加粗等），删除其余 run；
无 run 的段落新建 run；段落级样式（标题/对齐/缩进）与表格结构原样保留。
局限：嵌套表格暂按单元格段落平铺处理；图片/页眉页脚不参与文本替换。
"""
import json
import sys

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph


def extract(doc):
    blocks = []
    pi = 0
    ti = 0
    for child in doc.element.body.iterchildren():
        tag = child.tag
        if tag.endswith('}p'):
            p = Paragraph(child, doc)
            name = str(p.style.name or '')
            btype = 'heading' if (name.startswith('Heading') or name.startswith('标题')) else 'paragraph'
            blocks.append({'id': f'P{pi}', 'type': btype, 'text': p.text})
            pi += 1
        elif tag.endswith('}tbl'):
            tbl = Table(child, doc)
            for ri, row in enumerate(tbl.rows):
                for ci, cell in enumerate(row.cells):
                    for pi2, p in enumerate(cell.paragraphs):
                        blocks.append({
                            'id': f'T{ti}_R{ri}_C{ci}_P{pi2}',
                            'type': 'cell',
                            'text': p.text,
                        })
            ti += 1
    return blocks


def _set_text(p, text):
    runs = p.runs
    if runs:
        runs[0].text = text
        for r in runs[1:]:
            r._element.getparent().remove(r._element)
    else:
        p.add_run(text)


def main():
    mode = sys.argv[1]
    if mode == 'extract':
        doc = Document(sys.argv[2])
        print(json.dumps({'blocks': extract(doc)}, ensure_ascii=False))
    elif mode == 'apply':
        src, repl_path, out = sys.argv[2], sys.argv[3], sys.argv[4]
        repl = json.load(open(repl_path, encoding='utf-8'))
        doc = Document(src)
        by_id = {b['id']: b.get('text', '') for b in repl.get('blocks', [])}
        # 用与 extract 完全一致的顺序遍历，定位段落并替换
        pi = 0
        ti = 0
        applied = 0
        missing = []
        for child in doc.element.body.iterchildren():
            tag = child.tag
            if tag.endswith('}p'):
                p = Paragraph(child, doc)
                bid = f'P{pi}'
                pi += 1
                if bid in by_id and by_id[bid]:
                    _set_text(p, by_id[bid])
                    applied += 1
                elif bid not in by_id:
                    missing.append(bid)
            elif tag.endswith('}tbl'):
                tbl = Table(child, doc)
                for ri, row in enumerate(tbl.rows):
                    for ci, cell in enumerate(row.cells):
                        for pi2, p in enumerate(cell.paragraphs):
                            bid = f'T{ti}_R{ri}_C{ci}_P{pi2}'
                            if bid in by_id and by_id[bid]:
                                _set_text(p, by_id[bid])
                                applied += 1
                            elif bid not in by_id:
                                missing.append(bid)
                ti += 1
        doc.save(out)
        print(json.dumps({'applied': applied, 'missing': missing[:20]}, ensure_ascii=False))
    else:
        raise SystemExit('用法: docx_blocks.py extract|apply ...')


if __name__ == '__main__':
    main()
