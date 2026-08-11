# -*- coding: utf-8 -*-
"""
把 docs/competition/科技论文-SCULPTOR.md 渲染为带 Word 原生公式（OMML）的 docx。
用法：python3 scripts/gen-paper-docx.py
依赖：python-docx（公式为内嵌 OMML，无需 LaTeX 环境）
"""
import os, re
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import qn

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SRC = os.path.join(BASE, 'docs/competition/科技论文-SCULPTOR.md')
OUT = os.path.join(BASE, 'docs/competition/科技论文-SCULPTOR.docx')

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
M_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'


# ---------------- OMML（Word 原生公式）构建器 ----------------

def _esc(t):
    return t.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

def run(t):
    return f'<m:r><m:t xml:space="preserve">{_esc(t)}</m:t></m:r>'

def seq(*nodes):
    return ''.join(nodes)

def ssub(e, sub):
    return f'<m:sSub><m:e>{e}</m:e><m:sub>{sub}</m:sub></m:sSub>'

def ssup(e, sup):
    return f'<m:sSup><m:e>{e}</m:e><m:sup>{sup}</m:sup></m:sSup>'

def ssubsup(e, sub, sup):
    return f'<m:sSubSup><m:e>{e}</m:e><m:sub>{sub}</m:sub><m:sup>{sup}</m:sup></m:sSubSup>'

def frac(num, den):
    return f'<m:f><m:num>{num}</m:num><m:den>{den}</m:den></m:f>'

def rad(e):
    return (f'<m:rad><m:radPr><m:degHide m:val="1"/></m:radPr>'
            f'<m:deg/><m:e>{e}</m:e></m:rad>')

def nary(chr_, sub, sup, e):
    return (f'<m:nary><m:naryPr><m:chr m:val="{chr_}"/><m:limLoc m:val="undOvr"/></m:naryPr>'
            f'<m:sub>{sub}</m:sub><m:sup>{sup}</m:sup><m:e>{e}</m:e></m:nary>')

def delim(e, o='(', c=')'):
    return (f'<m:d><m:dPr><m:begChr m:val="{o}"/><m:endChr m:val="{c}"/></m:dPr>'
            f'<m:e>{e}</m:e></m:d>')

def acc(e):
    return (f'<m:acc><m:accPr><m:chr m:val="\u0302"/></m:accPr>'
            f'<m:e>{e}</m:e></m:acc>')


def build_equations():
    """论文四个公式的 OMML（式 1–4），序号随公式右置。"""
    vhat = acc(run('v'))
    vjA = ssubsup(acc(run('v')), run('j'), run('(A)'))
    vjB = ssubsup(acc(run('v')), run('j'), run('(B)'))
    vji = ssubsup(run('v'), run('j'), run('(i)'))
    s1, s2, s3, s4, s5 = (ssub(run('s'), run(str(k))) for k in range(1, 6))

    f1 = seq(
        vhat, run(' = '),
        delim(seq(run('句长均值'), run('，'), run('句长波动'), run('，'), run('短句占比'),
                  run('，'), run('口语度'), run('，'), run('意象密度'), run('，'),
                  run('情绪浓度'), run('，'), run('词汇丰富度'), run('，'),
                  run('语言新鲜度'))),
        run(' ∈ '), ssup(run('ℝ'), run('8')), run('　(1)'))

    f2 = seq(
        vjA, run(' = '),
        frac(delim(seq(vjA, run(' − '), ssub(run('min'), run('i')), vji)),
             delim(seq(ssub(run('max'), run('i')), vji))),
        run(' ∈ '), delim(seq(run('0'), run(','), run('1')), '[', ']'),
        run('　(2)'))

    f3 = seq(
        run('d'), delim(run('A,B')), run(' = '),
        rad(nary('∑', run('j=1'), run('8'),
                 ssup(delim(seq(vjA, run(' − '), vjB)), run('2')))),
        run('　(3)'))

    f4 = seq(
        run('C'), run(' = '), run('0.25'), run(' '), s1,
        run(' + '), run('0.20'), run(' '), s2,
        run(' + '), run('0.20'), run(' '), s3,
        run(' + '), run('0.20'), run(' '), s4,
        run(' + '), run('0.15'), run(' '), s5,
        run('　(4)'))
    return [f1, f2, f3, f4]


EQUATIONS = build_equations()


def add_equation(content):
    """把 OMML 公式作为块级元素插入文档主体（居中显示公式）。"""
    xml = (f'<m:oMathPara xmlns:m="{M_NS}" xmlns:w="{W_NS}">'
           f'<m:oMathParaPr><m:jc m:val="center"/></m:oMathParaPr>'
           f'<m:oMath>{content}</m:oMath></m:oMathPara>')
    doc.element.body.append(parse_xml(xml))


# ---------------- 文档排版 ----------------

doc = Document()
sec = doc.sections[0]
sec.left_margin = Cm(2.8); sec.right_margin = Cm(2.6)
sec.top_margin = Cm(2.8); sec.bottom_margin = Cm(2.6)


def set_font(run_, cn='宋体', en='Times New Roman', size=12, bold=False):
    run_.font.name = en
    run_.font.size = Pt(size)
    run_.font.bold = bold
    run_._element.rPr.rFonts.set(qn('w:eastAsia'), cn)


def para(text, cn='宋体', size=12, bold=False, align=None, indent=True, space=1.5):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    if indent:
        pf.first_line_indent = Pt(size * 2)
    pf.line_spacing = space
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_font(r, cn=cn, size=size, bold=bold)
    return p


def heading(text, level):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10 if level == 1 else 6)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.3
    r = p.add_run(text)
    set_font(r, cn='黑体' if level == 1 else '宋体',
             size={1: 16, 2: 14, 3: 12}[level], bold=True)


def add_table(rows):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            c = t.cell(i, j)
            c.text = ''
            r = c.paragraphs[0].add_run(cell)
            set_font(r, size=10.5, bold=(i == 0))


def add_image(path, caption=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_ = p.add_run()
    run_.add_picture(path, width=Cm(14.5))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cp.add_run(caption)
        set_font(r, size=10.5, bold=True)


MATH_KEYS = [
    (r'\hat{v} =', 0),
    (r'\frac{v_j', 1),
    (r'\sqrt{\sum', 2),
    ('0.25', 3),
]


def math_index(line):
    for key, idx in MATH_KEYS:
        if key in line:
            return idx
    return None


lines = open(SRC, encoding='utf-8').read().split('\n')
i = 0
table_buf = []
in_table = False
while i < len(lines):
    line = lines[i].rstrip()
    if not line.strip():
        i += 1
        continue
    if line.startswith('|'):
        if not in_table:
            in_table = True
            table_buf = []
        cells = [c.strip() for c in line.strip('|').split('|')]
        if not all(re.fullmatch(r':?-{2,}:?', c or '') for c in cells):
            table_buf.append(cells)
        i += 1
        continue
    if in_table:
        add_table(table_buf)
        in_table = False
        table_buf = []
        doc.add_paragraph()
    m = re.match(r'!\[(.*?)\]\((.*?)\)', line)
    if m:
        add_image(os.path.join(BASE, 'docs', 'competition',
                               os.path.basename(m.group(2))), caption=m.group(1))
        i += 1
        continue
    if line.startswith('$$') and line.endswith('$$') and len(line) > 4:
        idx = math_index(line)
        if idx is not None:
            add_equation(EQUATIONS[idx])
            i += 1
            continue
    if line.startswith('# '):
        para(line[2:].strip(), cn='黑体', size=22, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, space=1.2)
    elif line.startswith('## '):
        heading(line[3:].strip(), 1)
    elif line.startswith('### '):
        heading(line[4:].strip(), 2)
    elif line.startswith('#### '):
        heading(line[5:].strip(), 3)
    elif line.startswith('> '):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.3
        r = p.add_run(line[2:].strip())
        set_font(r, size=12)
    elif line.startswith('**关键词**'):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(line.strip())
        set_font(r, size=12, bold=True)
    elif re.match(r'^[-*] ', line):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.4
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Pt(18)
        r = p.add_run(line[2:].strip())
        set_font(r, size=12)
    else:
        para(line.strip())
    i += 1

if in_table and table_buf:
    add_table(table_buf)

doc.save(OUT)
print('saved:', OUT)
